"""
File Processing Module
Handles document upload, parsing, and chunking for VittoriaDB
"""

import os
import logging
import tempfile
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
import hashlib
import mimetypes

# Document processing imports
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of processed document"""
    id: str
    content: str
    metadata: Dict[str, Any]
    chunk_index: int
    total_chunks: int

@dataclass
class ProcessedDocument:
    """Represents a fully processed document"""
    id: str
    title: str
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    file_type: str

class FileProcessor:
    """Advanced file processing for RAG system"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize file processor"""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Supported file types
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html,
        }
    
    async def process_uploaded_file(self, 
                                   file_content: bytes,
                                   filename: str,
                                   additional_metadata: Dict[str, Any] = None) -> ProcessedDocument:
        """Process uploaded file and return structured document"""
        
        # Determine file type
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Generate document ID
        content_hash = hashlib.md5(file_content).hexdigest()[:8]
        doc_id = f"doc_{content_hash}_{filename.replace(' ', '_')}"
        
        # Create temporary file for processing
        with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Process file based on type
            processor = self.supported_types[file_ext]
            content, extracted_metadata = await processor(temp_file_path, filename)
            
            # Create base metadata
            base_metadata = {
                'filename': filename,
                'file_type': file_ext,
                'file_size': len(file_content),
                'content_hash': content_hash,
                'processing_timestamp': __import__('time').time(),
                **extracted_metadata
            }
            
            if additional_metadata:
                base_metadata.update(additional_metadata)
            
            # Create chunks
            chunks = self._create_chunks(content, doc_id, base_metadata)
            
            # Create processed document
            processed_doc = ProcessedDocument(
                id=doc_id,
                title=extracted_metadata.get('title', filename),
                content=content,
                chunks=chunks,
                metadata=base_metadata,
                file_type=file_ext
            )
            
            logger.info(f"✅ Processed {filename}: {len(chunks)} chunks, {len(content)} chars")
            return processed_doc
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_file_path)
            except:
                pass
    
    async def _process_pdf(self, file_path: str, original_filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'title': original_filename  # Use original filename as fallback
                }
                
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title') or original_filename,  # Use filename as fallback
                        'author': pdf_reader.metadata.get('/Author', 'Unknown'),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    })
                
                # Extract text
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                
                return text_content.strip(), metadata
                
        except Exception as e:
            logger.error(f"Failed to process PDF: {e}")
            raise ValueError(f"Failed to process PDF file: {str(e)}")
    
    async def _process_docx(self, file_path: str, original_filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or original_filename,  # Use original filename as fallback
                'author': doc.core_properties.author or 'Unknown',
                'subject': doc.core_properties.subject or '',
                'paragraphs': len(doc.paragraphs)
            }
            
            # Extract text
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            return text_content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    async def _process_text(self, file_path: str, original_filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                'title': original_filename,  # Use original filename directly
                'lines': len(content.split('\n')),
                'words': len(content.split())
            }
            
            return content, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                
                metadata = {
                    'title': os.path.basename(file_path),
                    'lines': len(content.split('\n')),
                    'words': len(content.split()),
                    'encoding': 'latin-1'
                }
                
                return content, metadata
            except Exception as e:
                raise ValueError(f"Failed to read text file: {str(e)}")
    
    async def _process_markdown(self, file_path: str, original_filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert to HTML then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text_content = soup.get_text()
            
            # Extract title from first heading
            title = original_filename  # Default to original filename
            first_heading = soup.find(['h1', 'h2', 'h3'])
            if first_heading:
                title = first_heading.get_text().strip()
            
            metadata = {
                'title': title,
                'format': 'markdown',
                'lines': len(md_content.split('\n')),
                'headings': len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']))
            }
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Failed to process Markdown: {e}")
            raise ValueError(f"Failed to process Markdown file: {str(e)}")
    
    async def _process_html(self, file_path: str, original_filename: str) -> Tuple[str, Dict[str, Any]]:
        """Process HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract title
            title_tag = soup.find('title')
            title = title_tag.get_text().strip() if title_tag else original_filename  # Use filename as fallback
            
            # Extract text content
            text_content = soup.get_text()
            
            metadata = {
                'title': title,
                'format': 'html',
                'links': len(soup.find_all('a')),
                'images': len(soup.find_all('img')),
                'headings': len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']))
            }
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Failed to process HTML: {e}")
            raise ValueError(f"Failed to process HTML file: {str(e)}")
    
    def _create_chunks(self, 
                      content: str, 
                      doc_id: str, 
                      base_metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Create overlapping chunks from document content"""
        
        if not content.strip():
            return []
        
        # Split content into sentences for better chunking
        sentences = self._split_into_sentences(content)
        chunks = []
        current_chunk = ""
        current_length = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence would exceed chunk size, create a chunk
            if current_length + sentence_length > self.chunk_size and current_chunk:
                chunk = self._create_chunk(
                    content=current_chunk.strip(),
                    doc_id=doc_id,
                    chunk_index=chunk_index,
                    base_metadata=base_metadata
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                current_chunk = overlap_text + " " + sentence
                current_length = len(current_chunk)
                chunk_index += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_length += sentence_length
        
        # Add final chunk if there's remaining content
        if current_chunk.strip():
            chunk = self._create_chunk(
                content=current_chunk.strip(),
                doc_id=doc_id,
                chunk_index=chunk_index,
                base_metadata=base_metadata
            )
            chunks.append(chunk)
        
        # Update total chunks count
        for chunk in chunks:
            chunk.total_chunks = len(chunks)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences for better chunking"""
        import re
        
        # Simple sentence splitting (can be improved with NLTK)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get overlap text from the end of current chunk"""
        if len(text) <= overlap_size:
            return text
        
        # Try to find a good breaking point (sentence boundary)
        overlap_text = text[-overlap_size:]
        sentence_start = overlap_text.find('. ')
        
        if sentence_start != -1:
            return overlap_text[sentence_start + 2:]
        
        return overlap_text
    
    def _create_chunk(self, 
                     content: str, 
                     doc_id: str, 
                     chunk_index: int,
                     base_metadata: Dict[str, Any]) -> DocumentChunk:
        """Create a document chunk with metadata"""
        
        chunk_id = f"{doc_id}_chunk_{chunk_index}"
        
        chunk_metadata = {
            **base_metadata,
            'chunk_index': chunk_index,
            'chunk_size': len(content),
            'content': content,  # Store content in metadata for search results
        }
        
        return DocumentChunk(
            id=chunk_id,
            content=content,
            metadata=chunk_metadata,
            chunk_index=chunk_index,
            total_chunks=0  # Will be updated later
        )

# Global file processor instance
_file_processor = None

def get_file_processor() -> FileProcessor:
    """Get or create global file processor instance"""
    global _file_processor
    if _file_processor is None:
        _file_processor = FileProcessor()
    return _file_processor
